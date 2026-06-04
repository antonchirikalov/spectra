#!/usr/bin/env python3
"""
build_word_form.py — generate interactive Word clarification form from JSON spec.

Native Word SDT controls (no macros, no protection required):
  - w14:checkbox  — clickable checkboxes (Word 2010+)
  - w:dropDownList — native dropdown lists (Word 2007+)
  - plain-text SDT fields with placeholder hint

Usage:
  python build_word_form.py --input form_spec.json --output output.docx
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree

# ── Namespaces ────────────────────────────────────────────────────────────────
W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W14 = "http://schemas.microsoft.com/office/word/2010/wordml"

# ── Brand palette (matches DTV clarifications style) ─────────────────────────
NAVY_HEX   = "061838"
BLUE2_HEX  = "1B4F8A"
BORDER_HEX = "C9D3DE"
SILVER_HEX = "F4F6F8"
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xC9, 0xD3, 0xDE)
TEXT_DARK  = RGBColor(0x1A, 0x1A, 0x2E)
GREY_LABEL = RGBColor(0x55, 0x55, 0x55)

_sdt_counter = 0


def _next_id() -> str:
    global _sdt_counter
    _sdt_counter += 1
    return str(_sdt_counter)


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _shade_cell(cell, fill_hex: str):
    """Set solid background fill on a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_borders(cell, color_hex: str = BORDER_HEX, sz: int = 6):
    """Apply uniform single border to a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color_hex)
        b.set(qn("w:space"), "0")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _no_borders(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _cell_padding(cell, top=4, bottom=4, left=6, right=6):
    """Set cell inner margin (in points * 20 = twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(int(val * 20)))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


# ── SDT (Structured Document Tag) builders ───────────────────────────────────

def _sdt_checkbox(checked: bool = False) -> etree._Element:
    """
    Native Word 2010+ checkbox SDT — clickable without protection.
    Uses w14:checkbox with Unicode ballot box characters.
    """
    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")
    lock = OxmlElement("w:lock")
    lock.set(qn("w:val"), "sdtContentLocked")

    cb_el = etree.SubElement(sdtPr, f"{{{W14}}}checkbox")
    chk   = etree.SubElement(cb_el, f"{{{W14}}}checked")
    chk.set(f"{{{W14}}}val", "1" if checked else "0")
    chk_state = etree.SubElement(cb_el, f"{{{W14}}}checkedState")
    chk_state.set(f"{{{W14}}}val", "2612")   # ☒ (U+2612)
    chk_state.set(f"{{{W14}}}font", "MS Gothic")
    unchk_state = etree.SubElement(cb_el, f"{{{W14}}}uncheckedState")
    unchk_state.set(f"{{{W14}}}val", "2610") # ☐ (U+2610)
    unchk_state.set(f"{{{W14}}}font", "MS Gothic")

    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"),   "MS Gothic")
    rFonts.set(qn("w:eastAsia"), "MS Gothic")
    rFonts.set(qn("w:hAnsi"),   "MS Gothic")
    rFonts.set(qn("w:hint"),    "eastAsia")
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")  # 11pt
    rPr.append(sz)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = "☒" if checked else "☐"
    r.append(t)
    sdtContent.append(r)
    sdt.append(sdtContent)
    return sdt


def _sdt_dropdown(options: list, default_idx: int = 0) -> etree._Element:
    """
    Native Word dropdown SDT — renders as clickable dropdown in Word.
    options: list of str.
    """
    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")
    ddl = OxmlElement("w:dropDownList")
    for i, opt in enumerate(options):
        item = OxmlElement("w:listItem")
        item.set(qn("w:displayText"), str(opt))
        item.set(qn("w:value"), str(i))
        ddl.append(item)
    sdtPr.append(ddl)
    sdt.append(sdtPr)

    display_text = options[default_idx] if options else "Choose..."

    sdtContent = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), "Normal")
    pPr.append(pStyle)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    pPr.append(jc)
    p.append(pPr)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "PlaceholderText")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = display_text
    r.append(t)
    p.append(r)
    sdtContent.append(p)
    sdt.append(sdtContent)
    return sdt


def _sdt_text(placeholder: str = "Click to enter text") -> etree._Element:
    """Plain-text SDT field with grey italic placeholder."""
    sdt = OxmlElement("w:sdt")

    sdtPr = OxmlElement("w:sdtPr")
    alias = OxmlElement("w:alias")
    alias.set(qn("w:val"), placeholder)
    sdtPr.append(alias)
    tag = OxmlElement("w:tag")
    tag.set(qn("w:val"), "text_input_" + _next_id())
    sdtPr.append(tag)
    showingPlcHdr = OxmlElement("w:showingPlcHdr")
    sdtPr.append(showingPlcHdr)
    sdt.append(sdtPr)

    sdtContent = OxmlElement("w:sdtContent")
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "AAAAAA")
    rPr.append(color)
    i_el = OxmlElement("w:i")
    rPr.append(i_el)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = placeholder
    r.append(t)
    p.append(r)
    sdtContent.append(p)
    sdt.append(sdtContent)
    return sdt


# ── Paragraph / run helpers ──────────────────────────────────────────────────

def _styled_run(para, text: str, bold=False, italic=False, size_pt=10,
                color: RGBColor | None = None) -> None:
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _add_para(container, text: str = "", bold=False, italic=False, size_pt=10,
              color: RGBColor | None = None,
              space_before=0, space_after=4,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    p = container.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        _styled_run(p, text, bold=bold, italic=italic, size_pt=size_pt, color=color)
    return p


def _inline_checkbox_para(container, label: str, checked: bool = False, size_pt=10):
    """Paragraph: [SDT checkbox]  <label>"""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    p._p.append(_sdt_checkbox(checked=checked))
    run = p.add_run(f"  {label}")
    run.font.size = Pt(size_pt)
    return p


def _inline_dropdown_para(container, label: str, options: list, size_pt=10):
    """Paragraph: <label>  [SDT dropdown]"""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if label:
        run = p.add_run(label + "  ")
        run.bold = True
        run.font.size = Pt(size_pt)
    p._p.append(_sdt_dropdown(options))
    return p


# ── Main form builder ─────────────────────────────────────────────────────────

class WordFormBuilder:

    def __init__(self, spec: dict):
        self.spec = spec
        self.doc  = Document()
        self._configure_page()

    # ── Document setup ────────────────────────────────────────────────────────

    def _configure_page(self):
        sec = self.doc.sections[0]
        sec.page_width    = Cm(21)
        sec.page_height   = Cm(29.7)
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)

    # ── Cover page ───────────────────────────────────────────────────────────

    def _build_cover(self):
        meta = self.spec.get("meta", {})
        doc  = self.doc

        # ── Navy header banner
        banner = doc.add_table(rows=1, cols=1)
        banner.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = banner.cell(0, 0)
        _shade_cell(cell, NAVY_HEX)
        _no_borders(cell)
        _cell_padding(cell, top=14, bottom=14, left=16, right=16)

        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(4)
        _styled_run(p1, meta.get("vendor", "Vendor"), bold=True, size_pt=22, color=WHITE)

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(4)
        _styled_run(p2, "BIDDER CLARIFICATION QUESTIONS", bold=True, size_pt=13, color=LIGHT_GREY)

        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        _styled_run(
            p3,
            f"{meta.get('project', '')}  |  {meta.get('reference', '')}  |  "
            f"Submission Round {meta.get('round', '1')}",
            size_pt=9, color=LIGHT_GREY,
        )

        doc.add_paragraph()

        # ── Meta table (Issued by / Submitted by / Date)
        mtbl = doc.add_table(rows=3, cols=2)
        mtbl.style = "Table Grid"
        mtbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for i, (k, v) in enumerate([
            ("Issued by",    meta.get("issued_by", "")),
            ("Submitted by", meta.get("vendor",    "")),
            ("Date",         meta.get("date",       datetime.today().strftime("%d %B %Y"))),
        ]):
            kc, vc = mtbl.cell(i, 0), mtbl.cell(i, 1)
            _shade_cell(kc, SILVER_HEX)
            _cell_padding(kc); _cell_padding(vc)
            kp = kc.paragraphs[0]
            _styled_run(kp, k, bold=True, size_pt=9)
            vp = vc.paragraphs[0]
            _styled_run(vp, v, size_pt=9)

        doc.add_paragraph()

        intro = meta.get("intro", "")
        if intro:
            ip = doc.add_paragraph(intro)
            ip.paragraph_format.space_after = Pt(6)
            ip.runs[0].font.size = Pt(9)

        doc.add_page_break()

    # ── Section header ────────────────────────────────────────────────────────

    def _section_header(self, title: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after  = Pt(8)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size  = Pt(13)
        run.font.color.rgb = RGBColor(0x06, 0x18, 0x38)

    # ── Question card ─────────────────────────────────────────────────────────

    def _build_card(self, q: dict):
        doc = self.doc

        q_id    = q.get("id", "Q-?")
        title   = q.get("title", "")
        context = q.get("context", "")
        default = q.get("default_assumption", "")
        owner   = q.get("owner_hint", "Project Manager")
        impact  = q.get("decision_impact", "")
        controls = q.get("controls", [])

        def _new_row(fill_hex: str | None = None):
            doc_tbl.add_row()
            c = doc_tbl.cell(len(doc_tbl.rows) - 1, 0)
            if fill_hex:
                _shade_cell(c, fill_hex)
            _set_borders(c)
            _cell_padding(c, top=6, bottom=6, left=10, right=10)
            return c

        # ── Outer single-column card table
        doc_tbl = doc.add_table(rows=1, cols=1)
        doc_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Row 0: Navy header with question ID + title
        hdr = doc_tbl.cell(0, 0)
        _shade_cell(hdr, NAVY_HEX)
        _set_borders(hdr)
        _cell_padding(hdr, top=8, bottom=8, left=10, right=10)
        hp = hdr.paragraphs[0]
        hp.paragraph_format.space_after = Pt(0)
        _styled_run(hp, f"{q_id}    ", bold=True, size_pt=11, color=LIGHT_GREY)
        _styled_run(hp, title,         bold=True, size_pt=11, color=WHITE)

        # Row: context paragraph
        if context:
            ctx = _new_row()
            cp = ctx.paragraphs[0]
            cp.paragraph_format.space_after = Pt(2)
            _styled_run(cp, context, size_pt=9, color=TEXT_DARK)

        # Row: default assumption (silver background)
        if default:
            def_cell = _new_row(SILVER_HEX)
            dp = def_cell.paragraphs[0]
            dp.paragraph_format.space_after = Pt(2)
            _styled_run(dp, "Default assumption (Bidder's working hypothesis)\n",
                        bold=True, size_pt=8, color=GREY_LABEL)
            _styled_run(dp, default, italic=True, size_pt=9, color=TEXT_DARK)

        # Row: confirm checkbox + owner field
        conf = _new_row()
        _inline_checkbox_para(conf, "Confirmed as stated", checked=False, size_pt=10)
        op = conf.add_paragraph()
        op.paragraph_format.space_before = Pt(4)
        op.paragraph_format.space_after  = Pt(4)
        _styled_run(op, "Owner / Best person to answer:  ", bold=True, size_pt=9)
        op._p.append(_sdt_text(placeholder=f"Name and role ({owner})"))

        # ── Override controls
        for ctrl in controls:
            ctype = ctrl.get("type", "")
            label = ctrl.get("label", "")

            if ctype == "dropdown":
                c = _new_row()
                _inline_dropdown_para(c, label, ctrl.get("options", ["Choose..."]))

            elif ctype == "checkbox_group":
                c = _new_row()
                lp = c.paragraphs[0]
                lp.paragraph_format.space_after = Pt(2)
                _styled_run(lp, label + ":", bold=True, size_pt=9)
                for item in ctrl.get("items", []):
                    _inline_checkbox_para(c, item.get("label", ""),
                                          checked=item.get("checked", False))

            elif ctype == "table":
                c = _new_row()
                if label:
                    lp = c.paragraphs[0]
                    lp.paragraph_format.space_after = Pt(4)
                    _styled_run(lp, label + ":", bold=True, size_pt=9)

                headers  = ctrl.get("headers", [])
                rows_data = ctrl.get("rows", [])
                n_cols   = len(headers)
                if n_cols > 0:
                    inner = c.add_table(rows=1 + len(rows_data), cols=n_cols)
                    inner.style = "Table Grid"
                    # Header row
                    for ci, hdr_txt in enumerate(headers):
                        hc = inner.cell(0, ci)
                        _shade_cell(hc, NAVY_HEX)
                        _cell_padding(hc, top=4, bottom=4, left=6, right=6)
                        hp2 = hc.paragraphs[0]
                        _styled_run(hp2, hdr_txt, bold=True, size_pt=8, color=WHITE)
                    # Data rows
                    for ri, row_spec in enumerate(rows_data):
                        row_label = row_spec.get("label", "")
                        fields    = row_spec.get("fields", [])
                        # First col: pre-filled label
                        lc = inner.cell(ri + 1, 0)
                        _shade_cell(lc, SILVER_HEX)
                        _cell_padding(lc, top=4, bottom=4, left=6, right=6)
                        _styled_run(lc.paragraphs[0], row_label, bold=True, size_pt=9)
                        # Remaining cols: SDT controls
                        for fi, field in enumerate(fields):
                            col_idx = fi + 1
                            if col_idx >= n_cols:
                                break
                            fc = inner.cell(ri + 1, col_idx)
                            _cell_padding(fc, top=4, bottom=4, left=6, right=6)
                            fp = fc.paragraphs[0]
                            ftype = field.get("type", "text")
                            if ftype == "dropdown":
                                fp._p.append(_sdt_dropdown(field.get("options", ["Choose..."])))
                            elif ftype == "checkbox":
                                fp._p.append(_sdt_checkbox(checked=False))
                                run = fp.add_run(f"  {field.get('label', '')}")
                                run.font.size = Pt(9)
                            else:  # text
                                fp._p.append(_sdt_text(field.get("placeholder", "")))

            elif ctype == "text_field":
                c = _new_row()
                lp = c.paragraphs[0]
                lp.paragraph_format.space_after = Pt(2)
                if label:
                    _styled_run(lp, label + "  ", bold=True, size_pt=9)
                lp._p.append(_sdt_text(ctrl.get("placeholder", "Enter text")))

        # Row: decision impact (silver footer)
        if impact:
            ic = _new_row(SILVER_HEX)
            ip2 = ic.paragraphs[0]
            ip2.paragraph_format.space_after = Pt(4)
            _styled_run(ip2, "Decision impact.  ", bold=True, size_pt=8, color=GREY_LABEL)
            _styled_run(ip2, impact, italic=True, size_pt=8, color=TEXT_DARK)

        # Spacing after card
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    # ── Assemble & save ──────────────────────────────────────────────────────

    def build(self) -> Document:
        self._build_cover()
        for section in self.spec.get("sections", []):
            self._section_header(section.get("title", ""))
            for q in section.get("questions", []):
                self._build_card(q)
        return self.doc

    def save(self, output_path: str):
        self.build().save(output_path)
        size_kb = Path(output_path).stat().st_size // 1024
        print(f"✅ Word form saved: {output_path} ({size_kb} KB)")


# ── CLI ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Build interactive Word clarification form from JSON spec"
    )
    parser.add_argument("--input",  required=True, help="Path to form_spec.json")
    parser.add_argument("--output", required=True, help="Path to output .docx")
    args = parser.parse_args()

    spec_path = Path(args.input)
    if not spec_path.exists():
        print(f"❌ Input not found: {spec_path}", file=sys.stderr)
        sys.exit(1)

    spec = json.loads(spec_path.read_text(encoding="utf-8"))
    WordFormBuilder(spec).save(args.output)


if __name__ == "__main__":
    main()
